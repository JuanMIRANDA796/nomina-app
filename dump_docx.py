
import sys
import os
from docx import Document

file_path = r"C:\Users\Juan\.gemini\antigravity\scratch\nomina_colombia\web\Taller_Proyecto\Modelo_Proyecto_UdeA_Entrega 5.docx"
output_path = r"C:\Users\Juan\.gemini\antigravity\scratch\nomina_colombia\web\Taller_Proyecto\extracted_text.txt"

try:
    doc = Document(file_path)
    with open(output_path, "w", encoding="utf-8") as f:
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(p.text + "\n")
        f.write("\n--- TABLES ---\n")
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")
    print(f"Text extracted to {output_path}")

except Exception as e:
    print(f"Error: {e}")
