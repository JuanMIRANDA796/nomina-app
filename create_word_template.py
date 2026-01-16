
import os
from docx import Document
from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

output_path = r"C:\Users\Juan\.gemini\antigravity\scratch\nomina_colombia\web\Taller_Proyecto\Analisis_Financiero_Template.docx"

doc = Document()

# Title
title = doc.add_heading('ANÁLISIS Y EVALUACIÓN FINANCIERA DEL PROYECTO', 0)
# title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("Este documento contiene la estructura para el análisis financiero solicitado. Por favor, complete cada sección con el texto correspondiente una vez tenga los datos del presupuesto.")

# Section 1
doc.add_heading('1. Análisis del Presupuesto', level=1)

doc.add_heading('1.1 Identificación de Rubros Mayores', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Identifique qué categorías (Marketing, Producción, etc.) consumen la mayor parte del presupuesto. Inserte aquí su análisis.]")

doc.add_heading('1.2 Coherencia Estratégica', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Justifique si la asignación de recursos es coherente con los objetivos del proyecto. ¿Se está invirtiendo lo suficiente en las áreas clave?] ")

doc.add_heading('1.3 Realismo y Sostenibilidad', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Analice si el presupuesto es realista y sostenible en el tiempo. ¿Los costos están subestimados? ¿Hay flujo de caja suficiente?]")

doc.add_heading('1.4 Optimización de Recursos', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Señale posibles áreas donde se podrían reducir costos o mejorar la eficiencia.]")


# Section 2
doc.add_heading('2. Evaluación Financiera (Indicadores)', level=1)
doc.add_paragraph("Si cuenta con proyecciones financieras detalladas, complete esta sección. De lo contrario, céntrese en el análisis presupuestal anterior.")

doc.add_heading('2.1 Indicadores de Rentabilidad (TIR, VPN)', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Inserte aquí el análisis de la Tasa Interna de Retorno (TIR) y el Valor Presente Neto (VPN).]")

doc.add_heading('2.2 Punto de Equilibrio', level=2)
doc.add_paragraph("[INSTRUCCIÓN: Analice el punto de equilibrio y si es alcanzable con las ventas proyectadas.]")

# Save
doc.save(output_path)
print(f"Word template created at: {output_path}")
