
import sys
import os

try:
    from docx import Document
    print("python-docx is installed.")
except ImportError:
    print("python-docx is NOT installed.")
    sys.exit(1)

file_path = r"C:\Users\Juan\.gemini\antigravity\scratch\nomina_colombia\web\Taller_Proyecto\Modelo_Proyecto_UdeA_Entrega 5.docx"

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    sys.exit(1)

try:
    doc = Document(file_path)
    print(f"Document loaded. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    
    print("\n--- Tables Content Preview ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows[:3]: # Print first 3 rows of each table
            cell_data = [cell.text.strip() for cell in row.cells]
            print(cell_data)

    print("\n--- Text Content Preview (First 500 chars) ---")
    full_text = "\n".join([p.text for p in doc.paragraphs])
    print(full_text[:500])

except Exception as e:
    print(f"Error processing docx: {e}")
